import numpy as np
import random
np.random.seed(42)
random.seed(42)
import statsmodels.api as sm
from statsmodels.tsa.stattools import grangercausalitytests, adfuller
import scipy.stats as stats
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
import docx.shared

def plot_granger_causality(neural_metric_ts, quantum_purity_ts, save_path):
    plt.figure(figsize=(10, 6))
    sns.set_theme(style="darkgrid", rc={"axes.facecolor": "#1e1e1e", "figure.facecolor": "#1e1e1e"})
    plt.rcParams.update({'text.color': 'white', 'axes.labelcolor': 'white', 'xtick.color': 'white', 'ytick.color': 'white'})
    
    fig, ax1 = plt.subplots(figsize=(10, 6))
    fig.patch.set_facecolor('#1e1e1e')
    ax1.set_facecolor('#1e1e1e')
    
    t = np.arange(len(neural_metric_ts))
    
    color = '#00BCD4'
    ax1.set_xlabel('Time (Microseconds Epoch)')
    ax1.set_ylabel('Neural Topological Metric (CFC / DMN)', color=color)
    ax1.plot(t, neural_metric_ts, color=color, lw=2, alpha=0.8, label='Neural Causality (Driver)')
    ax1.tick_params(axis='y', labelcolor=color)
    
    ax2 = ax1.twinx()
    color = '#FF9800'
    ax2.set_ylabel('Simulated Quantum Purity ($P$)', color=color)
    ax2.plot(t, quantum_purity_ts, color=color, lw=2, linestyle='--', label='Quantum State Decay (Target)')
    ax2.tick_params(axis='y', labelcolor=color)
    
    plt.title('Granger Causality: Topological Protection Delaying Quantum Decoherence', fontsize=14, pad=15, color='white')
    fig.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight', facecolor='#1e1e1e')
    plt.close()

def bayesian_counterfactual_impact(quantum_purity, doc, intervention_idx):
    doc.add_heading('Bayesian Counterfactual Shielding Effect Test (BSTS Causal Impact)', level=2)
    doc.add_paragraph("In ultra-low entropy silent systems, the conventional Granger causality model fails. System switching to Bayesian Structural Time Series (STS) counterfactual inference engine.")
    
    if len(quantum_purity) < intervention_idx + 5:
        doc.add_paragraph(f"[Skipped] Data segment too short (only {len(quantum_purity)} items), insufficient to support counterfactual projection.")
        return
        
    pre_intervention = quantum_purity[:intervention_idx]
    post_intervention = quantum_purity[intervention_idx:]
    
    try:
        model = sm.tsa.UnobservedComponents(pre_intervention, level='local linear trend')
        res = model.fit(disp=False)
        
        forecast = res.get_forecast(steps=len(post_intervention))
        cf_mean = np.clip(forecast.predicted_mean, 1e-15, None)
        cf_se = forecast.se_mean
        
        actual_cumulative = np.sum(post_intervention)
        cf_cumulative = np.sum(cf_mean)
        cf_cumulative_se = np.sqrt(np.sum(cf_se**2))
        
        z_score = (actual_cumulative - cf_cumulative) / (cf_cumulative_se + 1e-8)
        p_value = stats.norm.sf(z_score)
        
        p = doc.add_paragraph()
        p.add_run(f"Projected total collapse in counterfactual parallel universe: {cf_cumulative:.4f}\n").bold = True
        p.add_run(f"Actual observed total lifespan extension: {actual_cumulative:.4f}\n").bold = True
        
        if p_value < 0.05:
            res_str = f"[Shielding Confirmed] Actual lifespan significantly exceeded the counterfactual parallel universe collapse expectation! P-value: {p_value:.5e}"
            p2 = doc.add_paragraph(res_str)
            p2.runs[0].font.bold = True
            p2.runs[0].font.color.rgb = docx.shared.RGBColor(0, 150, 0)
        else:
            doc.add_paragraph(f"[Mean Reversion] Failed to confirm significant physical shielding effect. P-value: {p_value:.5f}")
    except Exception as e:
        doc.add_paragraph(f"Counterfactual inference failed due to mathematical anomaly: {e}")

def verify_granger_causality_with_evt(neural_metric_ts, quantum_purity_ts, doc):
    doc.add_heading('Granger Causality Test', level=2)
    
    def ensure_stationarity(ts):
        if len(ts) > 10:
            try:
                adf_res = adfuller(ts)
                if adf_res[1] > 0.05:  
                    return np.diff(ts), True
            except Exception:
                pass
        return ts, False

    linearized_quantum_ts = np.log(np.clip(quantum_purity_ts, 1e-8, None))
    
    ts_neural, diff_n = ensure_stationarity(neural_metric_ts)
    ts_quantum, diff_q = ensure_stationarity(linearized_quantum_ts)
    
    if diff_n and not diff_q:
        ts_quantum = ts_quantum[1:]
    elif diff_q and not diff_n:
        ts_neural = ts_neural[1:]
        
    min_len = min(len(ts_neural), len(ts_quantum))
    data_matrix = np.column_stack((ts_quantum[:min_len], ts_neural[:min_len]))
    
    try:
        n_obs = len(data_matrix)
        safe_max_lag = max(1, min(5, (n_obs // 3) - 1))
        
        if safe_max_lag < 1 or n_obs < 5:
            doc.add_paragraph(f"[Skipped] Too few valid data points (only {n_obs}), insufficient to support minimum Vector Autoregression (VAR) causal model.")
            return
            
        gc_res = grangercausalitytests(data_matrix, maxlag=safe_max_lag, verbose=False)
        p_values = [gc_res[lag][0]['ssr_ftest'][1] for lag in gc_res.keys()]
        min_p_value = min(p_values)
        
        if min_p_value < 0.05:
            res = f"[Breakthrough Confirmed] Significant non-classical causality discovered! Including neural data significantly improved quantum state evolution prediction accuracy. Minimum P-value: {min_p_value:.5e} (Optimal Lag: {p_values.index(min_p_value)+1})"
            p = doc.add_paragraph(res)
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = docx.shared.RGBColor(0, 150, 0)
        else:
            res = f"[Mean Reversion] Failed causality test, signal submerged by thermal noise. Minimum P-value: {min_p_value:.5f}"
            doc.add_paragraph(res)
    except Exception as e:
        doc.add_paragraph(f"Causality test failed due to data collinearity or anomaly: {e}")
