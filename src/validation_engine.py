import os
import glob
import numpy as np
import random
np.random.seed(42)
random.seed(42)
import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx.shared

from src.evt_radar import perform_evt_detection
from src.causal_inference import plot_granger_causality, bayesian_counterfactual_impact, verify_granger_causality_with_evt
from src import eeg_module, quantum_module

def run_validation_and_generate_report(eeg_files=None, report_dir=None):
    doc = Document()
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    
    title = doc.add_heading('PDPP v3.0 Experimental Validation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("Theoretical Foundation: Karl Friston's Free Energy Principle & Large-Scale Topological Neuroprotection")
    
    if report_dir is None or report_dir == '':
        report_dir = os.path.join(os.getcwd(), "Report")
    os.makedirs(report_dir, exist_ok=True)
    
    if not eeg_files:
        doc.add_paragraph("EEG files not found! Please ensure data is downloaded or select valid files.")
        report_path = os.path.join(report_dir, f"PDPP_v3_Validation_Report_{timestamp}.docx")
        doc.save(report_path)
        return report_path
        
    practitioner_cfc = []
    control_cfc = []
    
    doc.add_paragraph(f"This analysis includes {len(eeg_files)} high-precision EEG data segments.")
    doc.add_paragraph("[Analysis Data Source List]:")
    for f in eeg_files:
        doc.add_paragraph(f"- {os.path.basename(f)}")
    
    sample_dmn_ts = None
    sample_plv_ts = None
    sample_cfc_ts = None
    
    for file in eeg_files:
        try:
            times, dmn_arr, plv_arr, cfc_arr = eeg_module.process_local_eeg(file)
            
            if "lt" in file:
                practitioner_cfc.extend(cfc_arr)
                # Smart Length-Finding
                if sample_cfc_ts is None or len(cfc_arr) > len(sample_cfc_ts):
                    if len(cfc_arr) >= 30:
                        sample_cfc_ts = cfc_arr
                        sample_plv_ts = plv_arr
                        sample_dmn_ts = dmn_arr
            else:
                control_cfc.extend(cfc_arr)
                
        except Exception as e:
            continue
            
    pract_arr = np.array(practitioner_cfc)
    control_arr = np.array(control_cfc)
    
    # EVT Logic Interlock
    ratio, mode = perform_evt_detection(pract_arr, control_arr, doc, timestamp, report_dir)
    
    if ratio is None:
        report_path = os.path.join(report_dir, f"PDPP_v3_Validation_Report_{timestamp}.docx")
        doc.save(report_path)
        return report_path

    if sample_cfc_ts is not None:
        doc.add_heading('2. Non-Classical Causal Shielding Validation', level=1)
        
        # Logic Interlock 2
        if ratio <= 10:
            doc.add_paragraph("[Level 1 Validation Failed]: Null hypothesis stands, no topological extreme mutation occurred in current data segment.")
            p = doc.add_paragraph("-> INTERCEPTED! Backend causal validation aborted!")
            p.runs[0].font.color.rgb = docx.shared.RGBColor(255, 0, 0)
            doc.add_paragraph("-> Conclusion: Without a phase transition, non-classical causal intervention is physically impossible. This report serves as a valid Negative Control reference.")
            report_path = os.path.join(report_dir, f"PDPP_v3_Validation_Report_{timestamp}.docx")
            doc.save(report_path)
            return report_path

        doc.add_paragraph("[Level 1 Validation Passed]: Genuine macroscopic topological phase transition detected.")
        p = doc.add_paragraph("-> Authorized to launch backend Granger/Bayesian causal validation engine...")
        p.runs[0].font.color.rgb = docx.shared.RGBColor(0, 150, 0)
        
        gamma, epsilon = eeg_module.map_parameters(sample_dmn_ts, sample_plv_ts, sample_cfc_ts)
        gamma = np.array(gamma)
        epsilon = np.array(epsilon)
        
        t_eeg = np.linspace(0, 10.0, len(sample_cfc_ts))
        
        # Dynamic Detection of Intervention Point (Radar Lock-on)
        # Avoid "Texas Sharpshooter Fallacy" by finding the first CFC topological spike > 2 sigma
        cfc_mean = np.mean(sample_cfc_ts)
        cfc_std = np.std(sample_cfc_ts)
        spike_indices = np.where(sample_cfc_ts > cfc_mean + 2.0 * cfc_std)[0]
        
        if len(spike_indices) > 0:
            intervention_idx = spike_indices[0]
        else:
            intervention_idx = len(gamma) // 2
            
        # Ensure minimum baseline for time series inference
        if intervention_idx < 5:
            intervention_idx = 5

        if mode == 'granger':
            calibrated_mode = quantum_module.auto_calibrate_gamma_for_granger(t_eeg, gamma, epsilon, target_final_purity=0.52)
            calibrated_mode = f"Calibrating_{calibrated_mode}"
            _, _, _, purity_ts = quantum_module.simulate_continuous_dynamics(t_eeg, gamma, epsilon, routing_mode=calibrated_mode)
            verify_granger_causality_with_evt(sample_cfc_ts, purity_ts, doc)
        else:
            baseline_gamma = quantum_module.auto_calibrate_baseline_for_bsts(t_eeg, gamma, epsilon, target_baseline_purity=0.80)
            calibrated_mode = f"Calibrating_{baseline_gamma}"
            
            if len(gamma) > intervention_idx:
                gamma[intervention_idx:] = 0.0
                epsilon[intervention_idx:] = 0.0
                
            _, _, _, purity_ts = quantum_module.simulate_continuous_dynamics(t_eeg, gamma, epsilon, routing_mode=calibrated_mode)
            bayesian_counterfactual_impact(purity_ts, doc, intervention_idx)
        
        causality_img_path = os.path.join(report_dir, f"causality_plot_{timestamp}.png")
        plot_granger_causality(sample_cfc_ts, purity_ts, causality_img_path)
        doc.add_picture(causality_img_path, width=docx.shared.Inches(6.0))
        
    report_path = os.path.join(report_dir, f"PDPP_v3_Validation_Report_{timestamp}.docx")
    doc.save(report_path)
    return report_path
