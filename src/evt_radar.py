import numpy as np
import random
np.random.seed(42)
random.seed(42)
from scipy.stats import genextreme
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
import os
import docx.shared

def get_maxima(data, block_size):
    n_blocks = len(data) // block_size
    if n_blocks == 0: return np.array([np.max(data)])
    truncated = data[:n_blocks * block_size]
    blocks = truncated.reshape(n_blocks, block_size)
    return np.max(blocks, axis=1)

def plot_evt_distribution(max_pract, max_ctrl, shape_p, loc_p, scale_p, shape_c, loc_c, scale_c, threshold, save_path, mode="Burst_Phase_Transition"):
    plt.figure(figsize=(10, 6))
    sns.set_theme(style="darkgrid", rc={"axes.facecolor": "#1e1e1e", "figure.facecolor": "#1e1e1e"})
    plt.rcParams.update({'text.color': 'white', 'axes.labelcolor': 'white', 'xtick.color': 'white', 'ytick.color': 'white'})
    
    x = np.linspace(min(np.min(max_ctrl), np.min(max_pract)) * 0.8, max(np.max(max_ctrl), np.max(max_pract)) * 1.2, 500)
    
    pdf_p = genextreme.pdf(x, shape_p, loc=loc_p, scale=scale_p)
    pdf_c = genextreme.pdf(x, shape_c, loc=loc_c, scale=scale_c)
    
    plt.plot(x, pdf_c, color='#4CAF50', lw=2, label='Controls (Short-term)')
    plt.fill_between(x, pdf_c, alpha=0.3, color='#4CAF50')
    
    plt.plot(x, pdf_p, color='#E91E63', lw=2, label='Practitioners (Long-term)')
    plt.fill_between(x, pdf_p, alpha=0.3, color='#E91E63')
    
    plt.axvline(x=threshold, color='#FFC107', linestyle='--', lw=2, label=f'Extreme Threshold ({threshold:.3f})')
    
    x_tail = x[x >= threshold]
    pdf_p_tail = pdf_p[x >= threshold]
    
    if mode == "Silent_Phase_Transition":
        fill_label = 'Silent Phase Transition (Nirodha / Zero-Point)'
        plt.title('EVT: Left-Tail GEV Distribution (Inverted for Silent Collapse)', fontsize=14, pad=15)
        plt.xlabel('Negative Normalized PAC (Inverted ndPAC)')
    elif mode == "Classical_Noise":
        fill_label = 'Classical Thermal Noise Limit'
        plt.title('EVT: GEV Distribution (Classical Noise)', fontsize=14, pad=15)
        plt.xlabel('Normalized Phase-Amplitude Coupling (ndPAC)')
    else:
        fill_label = 'Topological Phase Transition (Freak Wave)'
        plt.xlabel('Normalized Phase-Amplitude Coupling (ndPAC)')
        
    plt.fill_between(x_tail, pdf_p_tail, alpha=0.8, color='#FFC107', label=fill_label)
    plt.ylabel('Probability Density')
    plt.legend(facecolor='#2d2d2d', edgecolor='none', labelcolor='white')
    
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight', facecolor='#1e1e1e')
    plt.close()

def perform_evt_detection(pract_arr, control_arr, doc, timestamp, report_dir):
    doc.add_heading('1. Extreme Value Theory (EVT) Topological Phase Transition Test', level=1)
    
    doc.add_paragraph(f"Long-term Practitioners Epochs extracted: {len(pract_arr)}")
    doc.add_paragraph(f"Short-term Controls Epochs extracted: {len(control_arr)}")
    
    if len(pract_arr) < 5 or len(control_arr) < 5:
        doc.add_paragraph("ERROR: Insufficient data volume for statistical EVT analysis.")
        return None, None
        
    block_size = max(2, len(control_arr) // 20)
    
    baseline_variance = np.var(control_arr)
    pract_variance = np.var(pract_arr)
    
    if baseline_variance < 1e-6 or pract_variance < 1e-6:
        p = doc.add_paragraph(f"[Error] Signal loss or zero variance (Pract: {pract_variance:.2e}, Ctrl: {baseline_variance:.2e}), suspected sensor failure or DC flatline.")
        p.runs[0].font.color.rgb = docx.shared.RGBColor(255, 0, 0)
        doc.add_paragraph("-> INTERCEPTED! System launch aborted, anomalous sample excluded.")
        return None, None
        
    variance_ratio = pract_variance / baseline_variance
    
    if variance_ratio > 1.2:
        mode = "Burst_Phase_Transition"
        sys_status = "System Judgment: High-frequency dynamic system (Generative Phase Transition) detected, initiating right-tail maxima breakthrough test."
        prob_title = "Probability of exceeding right-tail extremes (Freak Waves Probability):"
        calc_pract_arr = pract_arr
        calc_ctrl_arr = control_arr
    elif variance_ratio < 0.8:
        mode = "Silent_Phase_Transition"
        sys_status = "System Judgment: Ultra-low entropy silent system (Deconstructive Phase Transition) detected, initiating left-tail collapse test."
        prob_title = "Probability of breaching thermodynamic floor (Silent Collapse Probability):"
        calc_pract_arr = -pract_arr
        calc_ctrl_arr = -control_arr
    else:
        mode = "Classical_Noise"
        sys_status = "System Judgment: Data within conventional classical thermal noise range, maintaining mediocre state test."
        prob_title = "Incidence rate of maxima in classical thermal noise range (Classical Noise):"
        calc_pract_arr = pract_arr
        calc_ctrl_arr = control_arr
        
    p_status = doc.add_paragraph()
    p_status.add_run(sys_status).bold = True
    
    max_pract = get_maxima(calc_pract_arr, block_size)
    max_ctrl = get_maxima(calc_ctrl_arr, block_size)
    
    shape_p, loc_p, scale_p = genextreme.fit(max_pract)
    shape_c, loc_c, scale_c = genextreme.fit(max_ctrl)
    
    freq_mode = os.environ.get('PDPP_FREQ_MODE', 'high_gamma')
    if freq_mode == 'gamma':
        freq_label = "Theta [4-8Hz] Modulating Gamma [30-70Hz] ndPAC"
    else:
        freq_label = "Theta [4-8Hz] Modulating High-Gamma/Ripple [80-200Hz] ndPAC"
        
    doc.add_heading(f'Generalized Extreme Value (GEV) Fit Parameters ({freq_label})', level=2)
    doc.add_paragraph(f"Practitioners: Shape: {shape_p:.4f}, Loc: {loc_p:.4f}, Scale: {scale_p:.4f}")
    doc.add_paragraph(f"Controls: Shape: {shape_c:.4f}, Loc: {loc_c:.4f}, Scale: {scale_c:.4f}")
    
    extreme_threshold = np.percentile(np.concatenate([max_ctrl, max_pract]), 99.9) 
    prob_pract = genextreme.sf(extreme_threshold, shape_p, loc=loc_p, scale=scale_p)
    prob_ctrl = genextreme.sf(extreme_threshold, shape_c, loc=loc_c, scale=scale_c)
    
    p = doc.add_paragraph()
    p.add_run(prob_title).bold = True
    doc.add_paragraph(f"Configured coherence extreme threshold: {extreme_threshold:.4f} (ndPAC{' Inverted' if mode=='Silent_Phase_Transition' else ''})")
    doc.add_paragraph(f"Controls Probability: {prob_ctrl:.6e}")
    doc.add_paragraph(f"Practitioners Probability: {prob_pract:.6e}")
    
    if prob_ctrl == 0: prob_ctrl = 1e-15
    ratio = prob_pract / prob_ctrl
    
    if mode == "Silent_Phase_Transition":
        doc.add_paragraph(f"Probability Enhancement Ratio (Silent State / Control Thermal Noise) = {ratio:.2f}x")
    else:
        doc.add_paragraph(f"Probability Enhancement Ratio (Practitioners / Controls) = {ratio:.2f}x")
        
    doc.add_heading('Multiple Comparisons Correction (FDR)', level=2)
    import scipy.stats as stats
    from statsmodels.stats.multitest import multipletests
    
    # Empirical count of extreme events
    k_p = int(np.sum(max_pract >= extreme_threshold))
    k_c = int(np.sum(max_ctrl >= extreme_threshold))
    n_p = len(max_pract)
    n_c = len(max_ctrl)
    
    # Fisher's Exact Test for proportions
    table = [[k_p, n_p - k_p], [k_c, n_c - k_c]]
    _, raw_p_value = stats.fisher_exact(table, alternative='greater')
    
    # Apply Benjamini-Hochberg False Discovery Rate (FDR) correction
    # Architecture ready for array of p-values in massive multi-region testing
    reject, q_values, _, _ = multipletests([raw_p_value], alpha=0.05, method='fdr_bh')
    q_value = q_values[0]
    
    doc.add_paragraph(f"Raw Exact P-value (Fisher): {raw_p_value:.6e}")
    p_q = doc.add_paragraph(f"FDR-Adjusted Q-value (Benjamini-Hochberg): {q_value:.6e}")
    p_q.runs[0].font.bold = True
    
    doc.add_heading('Extreme Value Test Conclusion', level=2)
    if ratio > 10 and q_value < 0.05:
        doc.add_paragraph("Null hypothesis rejected (extremely high significance after FDR correction). Conclusion: Long-term meditators demonstrated extreme topological protection events of macroscopic quantum coherence. A probability difference of this magnitude cannot be explained by mean shifts in classical Gaussian white noise. This corroborates the pitchfork bifurcation model in PDPP theory.")
    else:
        doc.add_paragraph("Null hypothesis stands. Failed to capture significant extreme deviation phenomena after statistical correction.")

    evt_img_path = os.path.join(report_dir, f"evt_plot_{timestamp}.png")
    plot_evt_distribution(max_pract, max_ctrl, shape_p, loc_p, scale_p, shape_c, loc_c, scale_c, extreme_threshold, evt_img_path, mode=mode)
    doc.add_picture(evt_img_path, width=docx.shared.Inches(6.0))
    
    return ratio, mode
